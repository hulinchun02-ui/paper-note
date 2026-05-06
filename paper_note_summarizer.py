#!/usr/bin/env python3
"""Summarize academic PDFs into Chinese paper-note tables.

本文件是项目的主入口，既可以作为命令行程序运行，也可以被测试或
`skills/paper-note-summarizer/scripts/paper_note_summarizer.py` 包装脚本复用。

核心调用链：

1. `main()` 读取 `.env`、解析命令行参数、创建输出目录。
2. 非 dry-run 模式下，`extract_pdf_pages()` 从 PDF 中按页抽取文本。
3. `summarize_with_model()` 将页文本分块，调用 DeepSeek/OpenAI-compatible
   Chat Completions 接口生成结构化论文笔记。
4. `validate_and_normalize()` 校验模型 JSON，必要时由 `repair_summary_json()`
   再调用一次模型修复 schema。
5. `render_markdown()`、`render_docx()`、`write_evidence()` 将统一的 summary
   数据分别渲染为 Markdown、Word 表格和 evidence JSON。

所有渲染函数都依赖同一个 summary schema，因此新增字段时必须同步修改
`FIELD_SPECS`、prompt/schema 生成逻辑和对应测试。
"""

from __future__ import annotations

import argparse
import dataclasses
import json
import os
import re
from pathlib import Path
from typing import Any, Iterable


DEFAULT_BASE_URL = "https://api.deepseek.com"
DEFAULT_MODEL = "deepseek-v4-flash"
MISSING_TEXT = "论文未明确说明"
SCAN_TEXT_THRESHOLD = 500
CHUNK_CHAR_LIMIT = 26000


def load_env_file(env_path: Path = Path(".env")) -> None:
    """读取项目根目录 `.env` 中的 `KEY=VALUE` 配置。

    调用位置：`main()` 的第一步。

    设计原因：
    - 用户可以把 DeepSeek API key 写入 `.env`，不用每次手动 export。
    - 如果终端里已经设置了真实环境变量，则真实环境变量优先，`.env`
      不会覆盖它，方便临时切换模型或 key。
    - 这里只实现简单的 `KEY=VALUE` 解析，避免为一个小 CLI 增加额外依赖。
    """
    if not env_path.exists():
        return
    for raw_line in env_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


@dataclasses.dataclass(frozen=True)
class FieldSpec:
    """描述论文笔记中的一个输出字段。

    `FIELD_SPECS` 是整个项目的字段单一来源：
    - `schema_template()` 用它生成模型必须返回的 JSON schema。
    - `validate_and_normalize()` 用它校验模型输出是否完整。
    - `render_markdown()` 和 `render_docx()` 用它决定输出顺序和表格标签。

    字段含义：
    - `group_key` / `field_key`：机器可读的 JSON 路径。
    - `group_label` / `field_label`：用户可读的中文表格标题。
    - `long_text`：DOCX 渲染时是否按长文本单元格处理。
    """

    group_key: str
    group_label: str
    field_key: str
    field_label: str
    long_text: bool = False


FIELD_SPECS: tuple[FieldSpec, ...] = (
    FieldSpec("paper_overview", "论文概述", "title", "论文题目"),
    FieldSpec("paper_overview", "论文概述", "authors", "作者"),
    FieldSpec("paper_overview", "论文概述", "venue", "期刊/会议（如 ICLR2025）"),
    FieldSpec("paper_overview", "论文概述", "pdf_link", "论文 PDF 链接"),
    FieldSpec("paper_overview", "论文概述", "code_link", "论文代码（github 地址）"),
    FieldSpec("paper_content", "论文内容", "main_idea", "主要思想"),
    FieldSpec("paper_content", "论文内容", "problem", "所解决的问题"),
    FieldSpec(
        "paper_content",
        "论文内容",
        "main_content",
        "（1）论文的主要内容（这里介绍论文框架图以及框架图中的各个模块）",
        True,
    ),
    FieldSpec(
        "paper_content",
        "论文内容",
        "method_steps",
        "（2）文章所提方法的核心步骤（这里介绍推荐或者预测的主要流程）",
        True,
    ),
    FieldSpec(
        "paper_content",
        "论文内容",
        "additional_details",
        "其他补充（这里可以介绍方法部分其他细节，如其他重要图片的解释说明）",
        True,
    ),
    FieldSpec("experiments", "实验", "datasets", "数据集"),
    FieldSpec("experiments", "实验", "baselines", "对比方法"),
    FieldSpec("experiments", "实验", "metrics", "评价指标"),
    FieldSpec(
        "experiments",
        "实验",
        "results_conclusions",
        "实验结果及结论（这里介绍每个实验以及对应的结论）",
        True,
    ),
    FieldSpec("paper_highlights", "论文亮点", "highlights", "论文亮点", True),
    FieldSpec("paper_summary", "论文总结", "summary", "（个人理解，简要概括）", True),
)


def normalize_text(text: str) -> str:
    """清理 PDF 抽取出的页面文本。

    调用位置：`extract_pdf_pages()` 对每一页文本调用本函数。

    处理内容：
    - 去掉 PDF 中偶发的空字符。
    - 合并英文断词换行，例如 `anom-\naly` -> `anomaly`。
    - 压缩多余空白和连续空行。

    注意：这里不做语义改写，避免改变论文术语、公式或方法名。
    """
    text = text.replace("\x00", "")
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def extract_pdf_pages(pdf_path: Path) -> list[dict[str, Any]]:
    """从 PDF 中抽取按页组织的文本。

    输入：
    - `pdf_path`：本地论文 PDF 路径。

    输出：
    - `[{ "page_number": 1, "text": "..." }, ...]`

    调用位置：`main()` 在真实总结模式下调用，然后把结果传给
    `summarize_with_model()`。

    失败策略：
    - 缺少 `pypdf` 时给出安装提示。
    - 可抽取文本少于 `SCAN_TEXT_THRESHOLD` 时认为可能是扫描版 PDF，
      直接终止；当前版本不做 OCR。
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise SystemExit("缺少依赖 pypdf。请先运行：pip install -r requirements.txt") from exc

    reader = PdfReader(str(pdf_path))
    pages: list[dict[str, Any]] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = normalize_text(text)
        if text:
            pages.append({"page_number": index, "text": text})

    total_chars = sum(len(page["text"]) for page in pages)
    if total_chars < SCAN_TEXT_THRESHOLD:
        raise SystemExit("PDF 可抽取文本过少，可能是扫描版 PDF。v1 暂不做 OCR，请换成可复制文本的 PDF。")
    return pages


def pages_to_source(pages: Iterable[dict[str, Any]]) -> str:
    """把页对象拼接成带 `[Page N]` 标记的模型输入文本。

    调用位置：
    - 单 chunk 时由 `summarize_with_model()` 直接放入最终总结 prompt。
    - 多 chunk 时先放入 `chunk_extraction_prompt()` 做证据线索抽取。

    页码标记是 evidence 追溯的基础，模型会根据这些标记生成
    `evidence.page_refs`。
    """
    return "\n\n".join(f"[Page {page['page_number']}]\n{page['text']}" for page in pages)


def chunk_pages(pages: list[dict[str, Any]], limit: int = CHUNK_CHAR_LIMIT) -> list[list[dict[str, Any]]]:
    """按字符长度把论文页切成多个 chunk。

    调用位置：`summarize_with_model()`。

    为什么按页切：
    - 保留页码，便于 evidence 回溯。
    - 避免把单页拆开后模型难以定位上下文。

    `limit` 是粗略字符上限，不等同于 token 上限，但足以避免普通论文一次性
    塞入过长上下文。
    """
    chunks: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    current_len = 0
    for page in pages:
        page_len = len(page["text"]) + 32
        if current and current_len + page_len > limit:
            chunks.append(current)
            current = []
            current_len = 0
        current.append(page)
        current_len += page_len
    if current:
        chunks.append(current)
    return chunks


def schema_template() -> dict[str, Any]:
    """生成模型输出和本地校验共用的 summary JSON 模板。

    调用位置：
    - `final_summary_prompt()`：把模板嵌入 prompt，约束模型输出。
    - `validate_and_normalize()`：创建规范化结果的初始结构。
    - `repair_summary_json()`：告诉模型修复目标 schema。

    每个字段固定包含：
    - `value`：最终写入 DOCX/Markdown 的中文总结。
    - `evidence`：页码、短引文、可信度、备注，只写入 JSON 证据文件。
    """
    template: dict[str, Any] = {}
    for spec in FIELD_SPECS:
        template.setdefault(spec.group_key, {})
        template[spec.group_key][spec.field_key] = {
            "value": "",
            "evidence": {
                "page_refs": [],
                "source_quotes": [],
                "confidence": "low",
                "notes": "",
            },
        }
    return template


def system_prompt() -> str:
    """生成所有模型调用共享的系统提示词。

    调用位置：
    - `summarize_with_model()` 的分块证据抽取和最终总结调用。
    - `repair_summary_json()` 的 JSON 修复调用。

    核心约束是“只依据原文，不编造”，并要求模型输出合法 JSON。
    """
    return (
        "你是一名严谨的中文学术论文阅读助手。只能依据用户提供的论文原文总结，"
        "不得编造论文没有明确说明的信息。若没有证据，字段 value 必须写"
        f"“{MISSING_TEXT}”。所有输出必须是合法 JSON，不要输出 Markdown。"
    )


def final_summary_prompt(
    source_text: str,
    title_hint: str | None = None,
    pdf_link: str | None = None,
    code_link: str | None = None,
) -> str:
    """生成最终论文笔记总结 prompt。

    输入：
    - `source_text`：完整论文文本，或多 chunk 证据线索汇总。
    - `title_hint` / `pdf_link` / `code_link`：用户可选补充信息。

    输出：
    - 一个字符串 prompt，交给 `call_chat_completion()`。

    调用位置：`summarize_with_model()` 的最终总结阶段。

    这里会把 `schema_template()` 的 JSON 模板直接放进 prompt，降低模型漏字段
    或改字段名的概率。
    """
    template = json.dumps(schema_template(), ensure_ascii=False, indent=2)
    hints = {
        "title_hint": title_hint or "",
        "user_supplied_pdf_link": pdf_link or "",
        "user_supplied_code_link": code_link or "",
    }
    return f"""
请根据论文原文填写“论文笔记”JSON。

要求：
1. 严格保持 JSON 结构和字段名，不要新增顶层字段。
2. value 使用中文，学术严谨，优先概括方法、实验和结论的可验证事实。
3. evidence.page_refs 填相关页码整数列表。
4. evidence.source_quotes 只放短原文片段，每条不超过 40 个英文词或 80 个中文字符。
5. evidence.confidence 只能是 high、medium、low。
6. evidence.notes 说明证据不足、字段缺失或推断边界。
7. PDF/代码链接仅在用户提供或论文原文明确出现时填写，否则写“{MISSING_TEXT}”。
8. 实验结果及结论必须对应论文中的表格、指标、消融实验或文字结论。

用户补充信息：
{json.dumps(hints, ensure_ascii=False, indent=2)}

必须返回如下 JSON 结构：
{template}

论文原文：
{source_text}
""".strip()


def chunk_extraction_prompt(source_text: str) -> str:
    """生成长论文分块证据抽取 prompt。

    调用位置：`summarize_with_model()` 在 `chunk_pages()` 返回多个 chunk 时调用。

    这个阶段不生成最终论文笔记，只抽取 `topic/claim/page_refs/source_quotes`
    等线索。最终总结阶段再把各 chunk 的线索合并，减少长上下文压力。
    """
    return f"""
请从以下论文片段抽取后续总结可能需要的证据线索。只返回 JSON 对象，格式为：
{{"items": [{{"topic": "...", "claim": "...", "page_refs": [1], "source_quotes": ["..."]}}]}}
不要做最终总结，不要编造片段中没有的信息。

论文片段：
{source_text}
""".strip()


def call_chat_completion(
    messages: list[dict[str, str]],
    *,
    api_key: str,
    base_url: str,
    model: str,
    temperature: float = 0.1,
    timeout: int = 180,
) -> str:
    """调用 OpenAI-compatible Chat Completions 接口。

    输入：
    - `messages`：OpenAI chat 格式消息。
    - `api_key` / `base_url` / `model`：来自 CLI、环境变量或 `.env`。
    - `temperature`：默认较低，保证总结稳定。
    - `timeout`：HTTP 请求超时时间，单位秒。

    输出：
    - 模型返回的 `choices[0].message.content` 字符串。

    调用位置：
    - `summarize_with_model()`：分块证据抽取、最终总结。
    - `repair_summary_json()`：修复不合 schema 的 JSON。

    兼容策略：
    - 默认发送 `response_format={"type": "json_object"}`，要求 JSON。
    - 某些兼容服务不支持该参数时，如果 HTTP 400 且响应文本提到
      `response_format`，会移除该参数自动重试一次。
    """
    try:
        import requests
    except ImportError as exc:
        raise SystemExit("缺少依赖 requests。请先运行：pip install -r requirements.txt") from exc

    endpoint = base_url.rstrip("/") + "/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "response_format": {"type": "json_object"},
    }
    response = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)
    if response.status_code == 400 and "response_format" in response.text:
        payload.pop("response_format", None)
        response = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)
    if response.status_code >= 400:
        raise SystemExit(f"模型接口请求失败：HTTP {response.status_code}\n{response.text}")
    data = response.json()
    try:
        return data["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        preview = json.dumps(data, ensure_ascii=False)[:1000]
        raise SystemExit(f"模型接口返回格式异常：{preview}") from exc


def parse_json_object(text: str) -> dict[str, Any]:
    """从模型文本中解析 JSON 对象。

    调用位置：
    - `summarize_with_model()` 解析最终总结。
    - `repair_summary_json()` 解析修复后的总结。

    模型有时会把 JSON 包在 Markdown 代码块中，本函数会先剥离代码块；
    如果前后有解释文字，也会截取第一个 `{` 到最后一个 `}` 之间的内容。
    """
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    if not cleaned.startswith("{"):
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError("模型输出中未找到 JSON 对象")
        cleaned = cleaned[start : end + 1]
    return json.loads(cleaned)


def validate_and_normalize(data: dict[str, Any]) -> dict[str, Any]:
    """校验并规范化模型 summary JSON。

    输入：
    - 模型返回或 dry-run 生成的 summary 字典。

    输出：
    - 字段完整、类型稳定、顺序可由 `FIELD_SPECS` 控制的 summary 字典。

    调用位置：
    - `summarize_with_model()` 首次解析模型输出后调用。
    - `repair_summary_json()` 解析修复输出后调用。
    - `main()` 写文件前再调用一次，保证 dry-run 和真实路径一致。

    如果字段缺失或类型错误，会收集错误并抛出 `ValueError`。真实模型路径中，
    `summarize_with_model()` 会捕获该错误并交给 `repair_summary_json()`。
    """
    normalized = schema_template()
    errors: list[str] = []
    for spec in FIELD_SPECS:
        raw = data.get(spec.group_key, {}).get(spec.field_key)
        path = f"{spec.group_key}.{spec.field_key}"
        if not isinstance(raw, dict):
            errors.append(f"{path} must be an object")
            continue

        value = raw.get("value", MISSING_TEXT)
        if not isinstance(value, str) or not value.strip():
            errors.append(f"{path}.value must be a non-empty string")
            value = MISSING_TEXT

        evidence = raw.get("evidence", {})
        if not isinstance(evidence, dict):
            evidence = {}
            errors.append(f"{path}.evidence must be an object")

        page_refs = evidence.get("page_refs", [])
        if not isinstance(page_refs, list):
            page_refs = []
            errors.append(f"{path}.evidence.page_refs must be a list")
        page_refs = [int(page) for page in page_refs if isinstance(page, int) or str(page).isdigit()]

        source_quotes = evidence.get("source_quotes", [])
        if not isinstance(source_quotes, list):
            source_quotes = []
            errors.append(f"{path}.evidence.source_quotes must be a list")
        source_quotes = [str(quote).strip() for quote in source_quotes if str(quote).strip()]

        confidence = evidence.get("confidence", "low")
        if confidence not in {"high", "medium", "low"}:
            confidence = "low"
            errors.append(f"{path}.evidence.confidence must be high, medium, or low")

        notes = evidence.get("notes", "")
        normalized[spec.group_key][spec.field_key] = {
            "value": value.strip(),
            "evidence": {
                "page_refs": sorted(set(page_refs)),
                "source_quotes": source_quotes[:5],
                "confidence": confidence,
                "notes": notes.strip() if isinstance(notes, str) else str(notes),
            },
        }

    if errors:
        raise ValueError("; ".join(errors))
    return normalized


def repair_summary_json(
    invalid_json_text: str,
    validation_error: str,
    *,
    api_key: str,
    base_url: str,
    model: str,
) -> dict[str, Any]:
    """让模型修复不符合 schema 的 summary JSON。

    调用位置：`summarize_with_model()` 捕获解析或校验异常后调用。

    这里只修复结构，不重新阅读 PDF；prompt 会提供目标 schema、校验错误和
    原始模型输出，让模型尽量保留已生成内容并补齐缺失字段。
    """
    prompt = f"""
下面的 JSON 不符合目标 schema。请修复为合法 JSON，并严格使用目标字段。

校验错误：
{validation_error}

目标 schema：
{json.dumps(schema_template(), ensure_ascii=False, indent=2)}

待修复内容：
{invalid_json_text}
""".strip()
    repaired = call_chat_completion(
        [{"role": "system", "content": system_prompt()}, {"role": "user", "content": prompt}],
        api_key=api_key,
        base_url=base_url,
        model=model,
    )
    return validate_and_normalize(parse_json_object(repaired))


def summarize_with_model(
    pages: list[dict[str, Any]],
    *,
    api_key: str,
    base_url: str,
    model: str,
    title_hint: str | None,
    pdf_link: str | None,
    code_link: str | None,
) -> dict[str, Any]:
    """完成“论文文本 -> 结构化 summary”的核心模型流程。

    调用位置：`main()` 在 PDF 文本抽取完成后调用。

    流程：
    1. `chunk_pages()` 判断论文是否需要分块。
    2. 单 chunk：直接把带页码的原文送入 `final_summary_prompt()`。
    3. 多 chunk：每个 chunk 先用 `chunk_extraction_prompt()` 抽取证据线索，
       再把线索汇总送入最终总结 prompt。
    4. `call_chat_completion()` 调用 DeepSeek/OpenAI-compatible 接口。
    5. `parse_json_object()` + `validate_and_normalize()` 解析和校验。
    6. 如果失败，调用 `repair_summary_json()` 自动修复一次。

    输出会被 `main()` 继续传给 Markdown/DOCX/evidence 渲染函数。
    """
    chunks = chunk_pages(pages)
    if len(chunks) == 1:
        source_text = pages_to_source(chunks[0])
    else:
        evidence_notes = []
        for index, chunk in enumerate(chunks, start=1):
            content = call_chat_completion(
                [
                    {"role": "system", "content": system_prompt()},
                    {"role": "user", "content": chunk_extraction_prompt(pages_to_source(chunk))},
                ],
                api_key=api_key,
                base_url=base_url,
                model=model,
            )
            evidence_notes.append(f"[Chunk {index} evidence]\n{content}")
        source_text = "\n\n".join(evidence_notes)

    raw = call_chat_completion(
        [
            {"role": "system", "content": system_prompt()},
            {
                "role": "user",
                "content": final_summary_prompt(
                    source_text,
                    title_hint=title_hint,
                    pdf_link=pdf_link,
                    code_link=code_link,
                ),
            },
        ],
        api_key=api_key,
        base_url=base_url,
        model=model,
    )
    try:
        return validate_and_normalize(parse_json_object(raw))
    except Exception as exc:
        return repair_summary_json(raw, str(exc), api_key=api_key, base_url=base_url, model=model)


def sample_summary(title_hint: str | None = None) -> dict[str, Any]:
    """生成 dry-run 用的示例 summary。

    调用位置：`main()` 在 `--dry-run` 模式下调用。

    作用：
    - 不读取 PDF、不调用模型，也能验证 Markdown/DOCX/JSON 渲染链路。
    - 给测试用例提供稳定的输入数据。
    """
    data = schema_template()
    values = {
        ("paper_overview", "title"): title_hint or "Dry-run 示例论文",
        ("paper_overview", "authors"): "示例作者 A；示例作者 B",
        ("paper_overview", "venue"): "示例会议 2026",
        ("paper_overview", "pdf_link"): MISSING_TEXT,
        ("paper_overview", "code_link"): MISSING_TEXT,
        ("paper_content", "main_idea"): "该示例展示论文笔记工具的结构化输出格式。",
        ("paper_content", "problem"): "用于验证无模型调用时的渲染链路。",
        ("paper_content", "main_content"): "工具将论文信息拆分为论文概述、论文内容、实验、亮点和总结等部分。",
        ("paper_content", "method_steps"): "抽取 PDF 文本；调用模型生成结构化 JSON；校验 schema；渲染 DOCX、Markdown 和 evidence 文件。",
        ("paper_content", "additional_details"): "真实运行时此处应包含论文方法细节或关键图表解释。",
        ("experiments", "datasets"): "示例数据集",
        ("experiments", "baselines"): "示例对比方法",
        ("experiments", "metrics"): "示例评价指标",
        ("experiments", "results_conclusions"): "示例结果说明输出链路正常；真实运行时必须对应论文实验。",
        ("paper_highlights", "highlights"): "结构化、可核查、便于复用。",
        ("paper_summary", "summary"): "这是 dry-run 示例，不代表真实论文内容。",
    }
    for (group, field), value in values.items():
        data[group][field] = {
            "value": value,
            "evidence": {
                "page_refs": [1],
                "source_quotes": ["dry-run sample"],
                "confidence": "low",
                "notes": "这是 dry-run 样例数据，仅用于测试渲染。",
            },
        }
    return data


def value_of(summary: dict[str, Any], group: str, field: str) -> str:
    """从 summary 中取某个字段的 `value`。

    调用位置：
    - `render_markdown()` 写正文。
    - `render_docx()` 写表格右侧内容。

    这个小函数把 `summary[group][field]["value"]` 的访问集中起来，让渲染层
    不直接关心 evidence 结构。
    """
    return summary[group][field]["value"]


def render_markdown(summary: dict[str, Any], output_path: Path) -> None:
    """把 summary 渲染为 Markdown 论文笔记。

    调用位置：`main()` 在非 `--json-only` 模式下调用。

    Markdown 输出只写 `value`，不写 evidence；证据单独由 `write_evidence()`
    输出为 JSON，避免用户阅读版笔记过于拥挤。
    """
    lines = ["# 论文笔记", ""]
    current_group = None
    for spec in FIELD_SPECS:
        if spec.group_label != current_group:
            current_group = spec.group_label
            lines.extend([f"## {current_group}", ""])
        lines.extend([f"### {spec.field_label}", "", value_of(summary, spec.group_key, spec.field_key), ""])
    output_path.write_text("\n".join(lines), encoding="utf-8")


def set_cell_text(cell: Any, text: str, *, bold: bool = False, align_center: bool = False) -> None:
    """设置 DOCX 表格单元格文本和中文字体。

    调用位置：
    - `render_docx()` 写标题、左侧分组、字段标签、字段内容时反复调用。

    `python-docx` 对中文字体需要同时设置 `run.font.name` 和
    `w:eastAsia`，否则 Word 里可能回退到不一致的字体。
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(14 if bold else 11)


def set_table_borders(table: Any) -> None:
    """为 DOCX 表格设置黑色单线边框。

    调用位置：`render_docx()` 创建表格后立即调用。

    `python-docx` 没有高级边框 API，因此这里直接操作底层 OOXML。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def render_docx(summary: dict[str, Any], output_path: Path) -> None:
    """把 summary 渲染为截图样式的 DOCX 表格。

    调用位置：`main()` 在非 `--json-only` 模式下调用。

    渲染关系：
    - 第一行合并三列作为“论文笔记”标题。
    - 第 1 列按 `group_key` 合并为“论文概述/论文内容/实验/...”
      等大类。
    - 第 2 列写 `field_label`。
    - 第 3 列写 summary 中对应字段的 `value`。

    evidence 不放进 DOCX，保持表格简洁；需要核查时看 `.evidence.json`。
    """
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Cm
    except ImportError as exc:
        raise SystemExit("缺少依赖 python-docx。请先运行：pip install -r requirements.txt") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    set_table_borders(table)
    title_cell = table.rows[0].cells[0].merge(table.rows[0].cells[2])
    set_cell_text(title_cell, "论文笔记", bold=True, align_center=True)

    rows_by_group: dict[str, list[Any]] = {}
    for spec in FIELD_SPECS:
        row = table.add_row()
        rows_by_group.setdefault(spec.group_key, []).append(row)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(row.cells[1], spec.field_label, align_center=not spec.long_text)
        set_cell_text(row.cells[2], value_of(summary, spec.group_key, spec.field_key))

    for group_key, rows in rows_by_group.items():
        merged = rows[0].cells[0]
        for row in rows[1:]:
            merged = merged.merge(row.cells[0])
        group_label = next(spec.group_label for spec in FIELD_SPECS if spec.group_key == group_key)
        set_cell_text(merged, group_label, align_center=True)
        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(7.0)
        row.cells[2].width = Cm(16.0)

    document.save(output_path)


def write_evidence(summary: dict[str, Any], output_path: Path) -> None:
    """把 summary 中的 evidence 部分单独写成 JSON 文件。

    调用位置：`main()` 每次都会调用，不受 `--json-only` 影响。

    输出文件用于人工复核：每个字段保留页码、短原文片段、置信度和备注。
    """
    evidence: dict[str, Any] = {}
    for spec in FIELD_SPECS:
        evidence.setdefault(spec.group_key, {})
        evidence[spec.group_key][spec.field_key] = summary[spec.group_key][spec.field_key]["evidence"]
    output_path.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")


def safe_stem(pdf_path: Path | None, title_hint: str | None) -> str:
    """生成安全的输出文件名前缀。

    调用位置：`main()` 写四类输出文件前调用。

    优先使用 `title_hint`，否则使用 PDF 文件名；会把空格和特殊字符替换成
    下划线，避免中文/空格/标点导致跨平台路径问题。
    """
    raw = title_hint or (pdf_path.stem if pdf_path else "paper_note")
    stem = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", raw, flags=re.UNICODE).strip("._")
    return stem or "paper_note"


def build_arg_parser() -> argparse.ArgumentParser:
    """构建命令行参数解析器。

    调用位置：`main()` 读取 `.env` 后调用。

    重要接口：
    - `--pdf`：真实论文 PDF 路径。
    - `--out`：输出目录。
    - `--api-key` / `--base-url` / `--model`：覆盖 `.env` 或环境变量。
    - `--dry-run`：绕过 PDF 和模型，测试渲染链路。
    - `--json-only`：只生成 JSON，不生成 DOCX/Markdown。

    注意：默认值会读取 `DEEPSEEK_*`，再回退到 `OPENAI_*`，最后使用代码内
    默认 DeepSeek 配置。
    """
    parser = argparse.ArgumentParser(description="自动总结论文 PDF，生成 DOCX、Markdown 和 evidence JSON。")
    parser.add_argument("--pdf", type=Path, help="论文 PDF 路径。")
    parser.add_argument("--out", type=Path, default=Path("outputs"), help="输出目录。")
    parser.add_argument("--title-hint", help="可选论文题目提示，用于输出文件名和模型辅助。")
    parser.add_argument("--pdf-link", help="可选论文 PDF URL；若不提供且论文未写明，则输出“论文未明确说明”。")
    parser.add_argument("--code-link", help="可选代码仓库 URL；若不提供且论文未写明，则输出“论文未明确说明”。")
    parser.add_argument(
        "--api-key",
        default=os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENAI_API_KEY"),
        help="OpenAI 兼容 API key；优先读取 DEEPSEEK_API_KEY，其次 OPENAI_API_KEY。",
    )
    parser.add_argument(
        "--base-url",
        default=os.getenv("DEEPSEEK_BASE_URL") or os.getenv("OPENAI_BASE_URL", DEFAULT_BASE_URL),
        help="OpenAI 兼容 API base URL；默认 DeepSeek。",
    )
    parser.add_argument(
        "--model",
        default=os.getenv("DEEPSEEK_MODEL") or os.getenv("OPENAI_MODEL", DEFAULT_MODEL),
        help=f"模型名称；默认 {DEFAULT_MODEL}。",
    )
    parser.add_argument("--dry-run", action="store_true", help="不读取 PDF、不调用模型，生成示例输出以验证渲染链路。")
    parser.add_argument("--json-only", action="store_true", help="只写 summary JSON 和 evidence JSON，不生成 DOCX/MD。")
    return parser


def main(argv: list[str] | None = None) -> int:
    """命令行主流程。

    调用位置：
    - 文件被直接执行时，底部 `if __name__ == "__main__"` 调用。
    - skill 包装脚本通过 `runpy.run_path(..., run_name="__main__")` 间接调用。

    主流程负责把所有子模块串起来：
    1. `load_env_file()` 加载 `.env`。
    2. `build_arg_parser()` 解析 CLI。
    3. dry-run 走 `sample_summary()`；真实模式走
       `extract_pdf_pages()` + `summarize_with_model()`。
    4. `validate_and_normalize()` 做最终兜底校验。
    5. 写 `summary.json`、`evidence.json`，并按需写 Markdown 和 DOCX。
    """
    load_env_file()
    args = build_arg_parser().parse_args(argv)
    args.out.mkdir(parents=True, exist_ok=True)

    if args.dry_run:
        summary = sample_summary(args.title_hint)
        pdf_path = args.pdf
    else:
        if not args.pdf:
            raise SystemExit("请提供 --pdf，或使用 --dry-run 测试输出链路。")
        if not args.pdf.exists():
            raise SystemExit(f"PDF 文件不存在：{args.pdf}")
        if not args.api_key:
            raise SystemExit("缺少 API key。请设置 DEEPSEEK_API_KEY / OPENAI_API_KEY，或使用 --api-key。")
        pdf_path = args.pdf
        pages = extract_pdf_pages(args.pdf)
        summary = summarize_with_model(
            pages,
            api_key=args.api_key,
            base_url=args.base_url,
            model=args.model,
            title_hint=args.title_hint,
            pdf_link=args.pdf_link,
            code_link=args.code_link,
        )

    summary = validate_and_normalize(summary)
    stem = safe_stem(pdf_path, args.title_hint)
    summary_path = args.out / f"{stem}.summary.json"
    evidence_path = args.out / f"{stem}.evidence.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    write_evidence(summary, evidence_path)

    if not args.json_only:
        md_path = args.out / f"{stem}.md"
        docx_path = args.out / f"{stem}.docx"
        render_markdown(summary, md_path)
        render_docx(summary, docx_path)
        print(f"已生成：\n- {docx_path}\n- {md_path}\n- {evidence_path}\n- {summary_path}")
    else:
        print(f"已生成：\n- {evidence_path}\n- {summary_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
